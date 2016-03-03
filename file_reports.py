# -*- coding: utf-8 -*-

'''
    Helps to handle the mess in temp directory.
    Helps to create reports.
'''

import datetime
import os
import time

temp_dir = "temp\\"
logo_file = '__secret\\logo.png'

def _tmpname():
    # Уникальное имя для директории одного пользователя
    temp_name = datetime.datetime.now().strftime('%y%m%d_%H%M%S_%f')
    return temp_name

def clean_old_temp_files():
    today = datetime.datetime.now()
    for f_i in os.listdir(temp_dir):
        t = os.path.getmtime(temp_dir+f_i)
        days_wating = (today - datetime.datetime.fromtimestamp(t)).days
        if days_wating > 1:
            try:
                os.remove(temp_dir+f_i)
            except WindowsError:
                # Wish I use logging here one day...
                print('failed to cleanup ' + temp_dir+f_i)

def report_on_prices(agent=None):
    from pyexcel_ods import save_data
    from collections import OrderedDict
    import db_main
    data = OrderedDict()
    prices_list = db_main.get_prices_list(agent) # Спокойно принимает None
    print_list = []
    print_list.append([u"Материал",
                       u"Цена",
                       u"Валюта",
                       u"Условия платежа"])
    for pr_i in prices_list:
        if pr_i.is_for_group:
            mat_name = unicode(pr_i.material_type)
        else:
            mat_name = unicode(pr_i.material)
        print_list.append([mat_name,
                           pr_i.price_value,
                           pr_i.payterm.ccy_quote,
                           pr_i.payterm.payterm_name])
    data.update({'Prices': print_list})
    report_name = temp_dir+_tmpname() + "_prices.ods"
    save_data(report_name, data)
    os.startfile(report_name)

def report_on_all_client_prices():
    from pyexcel_ods import save_data
    from collections import OrderedDict
    import db_main
    data = OrderedDict()
    prices_list = db_main.get_sell_prices()
    print_list = []
    print_list.append([u"Материал",
                       u"Цена",
                       u"Валюта",
                       u"Условия платежа"])
    for pr_i in prices_list:
        if pr_i.is_for_group:
            mat_name = unicode(pr_i.material_type)
        else:
            mat_name = unicode(pr_i.material)
        print_list.append([mat_name,
                           pr_i.price_value,
                           pr_i.payterm.ccy_quote,
                           pr_i.payterm.payterm_name])
    data.update({'Prices': print_list})
    report_name = temp_dir+_tmpname() + "_prices.ods"
    save_data(report_name, data)
    os.startfile(report_name)

def report_on_all_supplier_prices():
    from pyexcel_ods import save_data
    from collections import OrderedDict
    import db_main
    data = OrderedDict()
    prices_list = db_main.get_buy_prices()
    print_list = []
    print_list.append([u"Материал",
                       u"Цена",
                       u"Валюта",
                       u"Incoterms",
                       u"Условия платежа",
                       u"Особые условия"])
    for pr_i in prices_list:
        spec_cond = u''
        if pr_i.is_only_for_sp_client:
            spec_cond += u' только для ' + unicode(pr_i.for_client)
        if pr_i.is_for_group:
            mat_name = unicode(pr_i.material_type)
        else:
            mat_name = unicode(pr_i.material)
        print_list.append([mat_name,
                           pr_i.price_value,
                           pr_i.payterm.ccy_quote,
                           pr_i.incoterms_place,
                           pr_i.payterm.payterm_name,
                           spec_cond])
    data.update({'Prices': print_list})
    report_name = temp_dir+_tmpname() + "_prices.ods"
    save_data(report_name, data)
    os.startfile(report_name)



def print_offer(agent_client):
    from docx import Document
    from docx.shared import Inches

    client_name = unicode(agent_client.name) #FIXME full_name
    to_persons = []
    for c_i in agent_client.contacts:
        if c_i.subs_to_prices and c_i.is_person:
            to_persons += [c_i.name]

    document = _get_company_docx_template(u'Коммерческое предложение', client_name, to_persons)
    # Now we have a template with logo. Print the letter itself

    document.add_paragraph(u'Предлагаем вашему вниманию:')




    # Save and open
    report_name = temp_dir + _tmpname() + "_offer.docx"
    document.save(report_name)
    os.startfile(report_name)

def _get_next_letter_number():
    n = datetime.datetime.now()
    let_number = n.strftime('%y%m%d_%H%M%S')
    let_date = n.strftime('%x')
    return [let_number, let_date]

def _get_company_docx_template(letter_header=u'', to_company=u'', to_persons=[]):
    '''
        Returns: docx.Document with logo, address e.t.c.
    '''
    from docx import Document
    from docx.shared import Inches

    document = Document()
    header_table = document.add_table(rows=1, cols=3)
    # Logo
    run = header_table.rows[0].cells[0].paragraphs[0].add_run()
    run.add_picture(logo_file, width=Inches(1.25))
    # Company details
    _write_company_details(header_table.rows[0].cells[1].paragraphs[0])
    # Client details
    run = header_table.rows[0].cells[2].paragraphs[0].add_run()
    run.add_text(to_company)
    run.add_break()
    if len(to_persons) > 0:
        run.add_text(u'Вниманию:')
    for p_i in to_persons:
        run.add_break()
        run.add_text(p_i)
    header_table.autofit = True

    n, d = _get_next_letter_number()
    document.add_paragraph(u'Исх. № ' + n + u' от ' + d)

    document.add_heading(letter_header, 1)

    return document

def _write_company_details(docx_paragraph):
    from cnf import get_cnf_text
    from docx.shared import Pt

    run = docx_paragraph.add_run()
    run.add_text(get_cnf_text("CompanyDataConfig", "company_name"))
    run.add_break()
    run.font.size = Pt(14)
    run.font.bold = True
    run = docx_paragraph.add_run()
    run.add_text(get_cnf_text("CompanyDataConfig", "company_name_eng"))
    run.add_break()
    run.add_text(get_cnf_text("CompanyDataConfig", "company_address"))
    run.add_break()
    run.add_text(u'тел/факс:' + get_cnf_text("CompanyDataConfig", "company_tel"))
    run.add_break()
    run.add_text(get_cnf_text("CompanyDataConfig", "company_email"))
    run.add_break()
    run.add_text(get_cnf_text("CompanyDataConfig", "company_tax_codes"))